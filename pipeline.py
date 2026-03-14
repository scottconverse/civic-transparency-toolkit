"""
Pipeline runner for Civic Transparency Toolkit.
Handles prompt loading, variable substitution, Claude API calls, and lane orchestration.
"""

import logging
import os
import re
import sys
import threading
import time
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)


def _strip_stage9_preamble(text: str) -> str:
    """Strip the research narration preamble from Stage 9 output.

    Stage 9 often starts with lines like 'I'll systematically research...'
    before the actual stories begin.  Stage 5 sees that preamble and
    misclassifies the whole output as 'not stories'.
    """
    # Strategy: find the first real content marker and strip everything before it
    markers = [
        re.search(r'^#{1,3}\s+(STORY|Story)\s', text, re.MULTILINE),
        re.search(r'^#{1,3}\s+\d+[.)]\s', text, re.MULTILINE),
        re.search(r'^#{1,3}\s+[A-Z][A-Z\s]{10,}', text, re.MULTILINE),  # ALL CAPS headline
        re.search(r'^---+\s*$', text, re.MULTILINE),  # --- separator line (story divider)
        re.search(r'^\*\*[A-Z]', text, re.MULTILINE),  # **Bold heading at line start
        re.search(r'^STORY\s+\d', text, re.MULTILINE),  # STORY 1: ...
        re.search(r'^CONFIDENCE RANKING', text, re.MULTILINE | re.IGNORECASE),
    ]
    # Find the earliest match
    first_story = None
    for m in markers:
        if m and (first_story is None or m.start() < first_story):
            first_story = m.start()

    if first_story and first_story > 100:
        # There's a real preamble (>100 chars of narration) — strip it
        stripped = text[first_story:]
        logger.info(f"Stripped {first_story} chars of preamble from Stage 9 output")
        return stripped
    elif first_story and first_story > 0:
        # Small preamble — still strip but log differently
        stripped = text[first_story:]
        logger.info(f"Stripped short preamble ({first_story} chars) from Stage 9 output")
        return stripped

    # Fallback: look for first double-newline followed by substantial text
    # (the preamble is usually 1-3 short sentences, then a blank line, then real content)
    paragraphs = text.split('\n\n')
    if len(paragraphs) > 1:
        # Check if first paragraph looks like narration (starts with "I'll", "I've", "Let me", etc.)
        first_para = paragraphs[0].strip()
        if re.match(r"^(I'll|I've|Let me|I will|I need to|Based on|I'm going|Here are|I can|I should)", first_para):
            rest = '\n\n'.join(paragraphs[1:])
            logger.info(f"Stripped narration preamble ({len(first_para)} chars): '{first_para[:60]}...'")
            return rest

    return text


def _extract_topic_briefs(text: str, city_name: str = "") -> str:
    """Extract a compact topic list from Stage 4 output for Stage 9.

    Stage 4 (Dark Signal Desk) produces ~50KB of adversarial analysis with
    triage labels, confidence assessments, and editorial conclusions that
    cause Stage 9 to refuse to write. Instead of passing the full analysis,
    we extract just the signal names, a brief description, and source URLs
    into a clean topic brief that Stage 9 can research independently.
    """
    if not text:
        return text

    # Extract signal headers: SIGNAL [S-1]: NAME or SIGNAL S-1: NAME
    signals = re.findall(
        r'SIGNAL\s*\[?S-(\d+)\]?:?\s*(.+?)(?:\n|$)', text
    )
    # Deduplicate (same signal may appear multiple times in the report)
    seen = set()
    unique_signals = []
    for num, name in signals:
        if num not in seen:
            seen.add(num)
            # Clean up markdown bold, emojis, and triage labels from name
            clean_name = name.strip().rstrip('*').strip()
            clean_name = re.sub(r'[⚠️🔴🟡🟢❌✓✗]+', '', clean_name).strip()
            clean_name = re.sub(r'\s*(CRITICAL|ESCALATE|HOLD|DISCARD)\s*$',
                                '', clean_name, flags=re.IGNORECASE).strip()
            unique_signals.append((num, clean_name))

    # Extract all URLs from the full text
    urls = list(set(re.findall(r'https?://[^\s\)\"\'<>]+', text)))

    if not unique_signals:
        # Fallback: if we can't parse signal structure, pass through
        # a truncated version without the most problematic sections
        logger.warning("Could not extract signals from Stage 4 output; "
                       "passing truncated text")
        # At least strip the worst offenders
        text = re.sub(r'(?i)(DISCARD|HOLD FOR PATTERN|FOR VERIFICATION'
                       r'|HANDOFF|ESCALAT)', '[--]', text)
        return text[:8000]  # Hard cap

    # Build clean topic brief
    lines = [
        "TOPIC LEADS FOR INDEPENDENT RESEARCH",
        "=" * 40,
        "",
        f"The following {len(unique_signals)} topics were identified in "
        f"{city_name + ' ' if city_name else ''}civic activity. Research each one independently.",
        "",
    ]
    for num, name in unique_signals:
        lines.append(f"TOPIC {num}: {name}")
        lines.append("")

    if urls:
        lines.append("")
        lines.append("REFERENCE URLS (starting points for research):")
        for url in sorted(urls):
            lines.append(f"- {url}")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Pipeline lane definitions
# ---------------------------------------------------------------------------

LANES = {
    "daily_production": {
        "name": "Find Today's News",
        "description": (
            "The everyday workflow. Scans all of your City & Sources for current news — "
            "council agendas, school board actions, local reporting, YouTube meeting transcripts — "
            "then writes full stories, fact-checks them, and rewrites everything in plain language "
            "your community can share. No input needed; just press Go."
        ),
        "stages": [1, 2, 5, 7],
    },
    "signal_intelligence": {
        "name": "Dig Deeper",
        "description": (
            "Looks for hidden patterns and connections your regular news sources aren't covering. "
            "Reads between the lines of public records, community forums, and meeting transcripts "
            "to find leads worth investigating. The output is investigative signals — questions "
            "worth asking — not finished stories. Press Go to scan your sources."
        ),
        "stages": [3, 4],
    },
    "investigate_and_write": {
        "name": "Investigate Hidden Stories",
        "description": (
            "The full investigation pipeline. Finds hidden patterns in your sources, then "
            "independently researches each one and writes complete, sourced stories — followed "
            "by a fact-check that catches errors before anything reaches your community. "
            "This is the most thorough workflow and produces the most detailed output. "
            "No input needed; press Go to start."
        ),
        "stages": [3, 9, 5],
    },
    "adhoc_story": {
        "name": "Write a Story",
        "description": (
            "Type a topic in the field above and the toolkit will research it from scratch — "
            "searching public records and news sources — then write a complete story with "
            "source citations, followed by a fact-check. Be specific: 'What's happening with "
            "the new housing development on Hover Street?' works better than 'Tell me about housing.'"
        ),
        "stages": [9, 5],
    },
    "story_polish": {
        "name": "Polish Existing Stories",
        "description": (
            "Already have story drafts? Paste them into the text box above. The toolkit will "
            "expand them with additional context and sourcing, rewrite them in plain shareable "
            "language, and run a fact-check. Good for cleaning up output from a previous run "
            "or polishing your own writing."
        ),
        "stages": [2, 7, 5],
    },
    "single_prompt": {
        "name": "Run a Single Step",
        "description": (
            "Run any one of the nine tools by itself. Pick a step from the list — the toolkit "
            "will show you what input it needs (if any) and what it does. Useful for running "
            "the fact-checker on a story you wrote yourself, looking up a legal question with "
            "First Amendment Counsel, or experimenting with individual tools."
        ),
        "stages": [],  # User picks which prompt
    },
}

PROMPT_NAMES = {
    1: "News Aggregator",
    2: "Story Expansion",
    3: "Black Desk — Speculative Radar",
    4: "Dark Signal Desk — Verification",
    5: "Integrity Checker",
    6: "First Amendment Counsel",
    7: "Community News Writer",
    8: "Source Integrity Protocol",
    9: "Story Research & Writing",
}

# One-line descriptions shown when picking a single step to run
PROMPT_DESCRIPTIONS = {
    1: "Scans your City & Sources for today's news, meetings, and public records.",
    2: "Expands the leads from Step 1 into full draft stories.",
    3: "Finds hidden patterns & connections using your City & Sources.",
    4: "Stress-tests the leads from Step 3 — separates real signals from noise.",
    5: "Fact-checks a story for accuracy, sourcing, and fairness.",
    6: "Checks a story for First Amendment and public records law issues.",
    7: "Rewrites a story in plain language your neighbors can share.",
    8: "Audits a story's sources for reliability and bias.",
    9: "You pick the topic — it researches and writes the story from scratch.",
}

PROMPT_FILES = {
    1: "01-news-aggregator.md",
    2: "02-story-expansion.md",
    3: "03-black-desk.md",
    4: "04-dark-signal-desk.md",
    5: "05-integrity-checker.md",
    6: "06-first-amendment-counsel.md",
    7: "07-plain-language-translator.md",
    8: "08-civic-grounding.md",
    9: "09-story-research-writing.md",
}


def get_prompts_dir():
    """Return the path to the prompts directory."""
    if getattr(sys, "frozen", False):
        base = Path(sys.executable).parent
        # PyInstaller --onedir puts add-data files in _internal/
        prompts = base / "_internal" / "prompts"
        if prompts.exists():
            return prompts
        return base / "prompts"
    else:
        return Path(__file__).parent / "prompts"


def load_prompt_text(prompt_number):
    """Load a prompt file and extract the content between the ``` markers."""
    prompts_dir = get_prompts_dir()
    filename = PROMPT_FILES.get(prompt_number)
    if not filename:
        raise ValueError(f"Unknown prompt number: {prompt_number}")

    filepath = prompts_dir / filename
    if not filepath.exists():
        raise FileNotFoundError(f"Prompt file not found: {filepath}")

    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    # Extract content between ``` markers (the actual prompt body).
    # Some prompts have multiple code blocks — concatenate them all.
    blocks = re.findall(r"```\n(.*?)```", content, re.DOTALL)
    if blocks:
        return "\n\n".join(b.strip() for b in blocks)
    # Fallback: return full content minus the header markdown
    return content


def substitute_variables(text, config):
    """Replace all [BRACKETED] template variables with values from config."""
    # Helper: return a non-empty string or the original bracket placeholder.
    # Guards against None and empty-string values leaking into prompts.
    def _v(val, fallback):
        return val if val else fallback

    city = _v(config.get("city_name"), "[YOUR_CITY]")
    state = _v(config.get("state"), "[YOUR_STATE]")
    sources = config.get("sources") or {}

    # Derive subreddit from city name (lowercase, no spaces) for Prompt 03
    city_sub = config.get("city_name", "")
    city_subreddit = city_sub.replace(" ", "") if city_sub else "[YOUR_CITY_SUBREDDIT]"

    replacements = {
        "[YOUR_CITY]": city,
        "[YOUR_STATE]": state,
        "[YOUR_CITY_AGENDA_PORTAL_URL]": _v(sources.get("city_agenda_portal_url"), "[YOUR_CITY_AGENDA_PORTAL_URL]"),
        "[YOUR_CITY_NEWS_URL]": _v(sources.get("city_news_url"), "[YOUR_CITY_NEWS_URL]"),
        "[YOUR_SCHOOL_DISTRICT]": _v(sources.get("school_district_name"), "[YOUR_SCHOOL_DISTRICT]"),
        "[YOUR_SCHOOL_DISTRICT_URL]": _v(sources.get("school_district_url"), "[YOUR_SCHOOL_DISTRICT_URL]"),
        "[YOUR_TRANSIT_AGENCY]": _v(sources.get("transit_agency_name"), "[YOUR_TRANSIT_AGENCY]"),
        "[YOUR_TRANSIT_AGENCY_URL]": _v(sources.get("transit_agency_url"), "[YOUR_TRANSIT_AGENCY_URL]"),
        "[YOUR_COUNTY]": _v(sources.get("county_name"), "[YOUR_COUNTY]"),
        "[YOUR_COUNTY_GOV_URL]": _v(sources.get("county_gov_url"), "[YOUR_COUNTY_GOV_URL]"),
        "[YOUR_STATE_GOV_URL]": _v(sources.get("state_gov_url"), "[YOUR_STATE_GOV_URL]"),
        "[YOUR_LOCAL_MEDIA_OUTLET]": _v(sources.get("local_media_outlet"), "[YOUR_LOCAL_MEDIA_OUTLET]"),
        "[YOUR_LOCAL_MEDIA_URL]": _v(sources.get("local_media_url"), "[YOUR_LOCAL_MEDIA_URL]"),
        "[YOUR_YOUTUBE_CHANNEL_URL]": _v(sources.get("youtube_channel_url"), "[YOUR_YOUTUBE_CHANNEL_URL]"),
        "[YOUR_STATE_OPEN_RECORDS_LAW]": _v(config.get("state_open_records_law"), "[YOUR_STATE_OPEN_RECORDS_LAW]"),
        "[YOUR_CITY_CLERK_CONTACT]": _v(config.get("city_clerk_contact"), "[YOUR_CITY_CLERK_CONTACT]"),
        "[YOUR_CITY_SUBREDDIT]": city_subreddit,
        "[DATE]": datetime.now().strftime("%B %d, %Y"),
    }

    # Composite replacement — only when BOTH city and state are real values
    # (not bracket placeholders). Must run BEFORE individual replacements
    # so the composite token is still intact.
    if city != "[YOUR_CITY]" and state != "[YOUR_STATE]":
        text = text.replace("[YOUR_CITY], [YOUR_STATE]", f"{city}, {state}")

    for placeholder, value in replacements.items():
        if value:
            text = text.replace(placeholder, value)

    return text


def build_user_message(prompt_number, previous_output=None, topic=None,
                       config=None, transcript_context=""):
    """Build the user message to send with each prompt."""
    if config is None:
        config = {}
    city = config.get("city_name", "your city") or "your city"
    today = datetime.now().strftime("%A, %B %d, %Y")

    # Build additional sources context if available
    additional_sources = ""
    if config and config.get("additional_sources"):
        source_lines = []
        tier_labels = {"A": "Official Record", "B": "News", "C": "Community"}
        for s in config["additional_sources"]:
            name = s.get("name", "Unnamed")
            url = s.get("url", "")
            tier = s.get("tier", "B")
            cat = s.get("category", "General")
            label = tier_labels.get(tier, "General")
            source_lines.append(f"- {name} | {url} | {label} | {cat}")
        if source_lines:
            additional_sources = (
                "\n\nADDITIONAL SOURCES TO CHECK:\n" + "\n".join(source_lines)
            )

    if prompt_number == 1:
        yt_url = config.get("sources", {}).get("youtube_channel_url", "")
        yt_line = f"- City YouTube channel: {yt_url}\n" if yt_url else ""

        # Build meeting schedule context if configured
        meeting_ctx = ""
        if config:
            sched = config.get("meeting_schedule", {})
            sched_lines = []
            if sched.get("council_day"):
                t = sched.get("council_time", "")
                sched_lines.append(f"City Council meets: {sched['council_day']}"
                                   + (f" at {t}" if t else ""))
            if sched.get("school_board_day"):
                t = sched.get("school_board_time", "")
                sched_lines.append(f"School Board meets: {sched['school_board_day']}"
                                   + (f" at {t}" if t else ""))
            if sched.get("county_commissioners_day"):
                t = sched.get("county_commissioners_time", "")
                sched_lines.append(f"County Commissioners meet: {sched['county_commissioners_day']}"
                                   + (f" at {t}" if t else ""))
            if sched.get("planning_commission_day"):
                t = sched.get("planning_commission_time", "")
                sched_lines.append(f"Planning Commission meets: {sched['planning_commission_day']}"
                                   + (f" at {t}" if t else ""))
            if sched_lines:
                meeting_ctx = (
                    "\n\nKNOWN MEETING SCHEDULE (check for recent agendas/minutes):\n"
                    + "\n".join(f"- {line}" for line in sched_lines)
                )

        return (
            f"Run the daily civic news aggregation for {city} for today, {today}. "
            f"You MUST use web_search to access REAL, CURRENT data. Do NOT simulate or fabricate results. "
            f"Search these sources for new developments, story leads, and civic intelligence:\n"
            f"- City agenda portal: {config.get('sources', {}).get('city_agenda_portal_url', 'N/A')}\n"
            f"- City news page: {config.get('sources', {}).get('city_news_url', 'N/A')}\n"
            f"- School district: {config.get('sources', {}).get('school_district_url', 'N/A')}\n"
            f"- County government: {config.get('sources', {}).get('county_gov_url', 'N/A')}\n"
            f"- Local media: {config.get('sources', {}).get('local_media_url', 'N/A')}\n"
            f"- Transit agency: {config.get('sources', {}).get('transit_agency_url', 'N/A')}\n"
            f"{yt_line}"
            f"Search each source individually. If a source returns no results, note that and move on. "
            f"Only report facts you can verify through web search.{additional_sources}"
            f"{meeting_ctx}"
            f"{transcript_context}"
        )
    elif prompt_number == 2:
        return (
            f"Here are the story leads from today's aggregation for {city}. "
            f"Select the strongest leads and expand them into community-ready stories.\n\n"
            f"--- BEGIN AGGREGATION OUTPUT ---\n{previous_output or '[No aggregation output provided]'}\n--- END AGGREGATION OUTPUT ---"
            f"{transcript_context}"
        )
    elif prompt_number == 3:
        return (
            f"Run the Black Desk speculative signal radar for {city} as of {today}. "
            f"You MUST use web_search to find REAL data. Do NOT simulate or fabricate signals. "
            f"Execute all three detection postures: anomaly scan, whisper correlation, "
            f"and bifurcation audit. Search local news, government sites, and public records "
            f"for actual anomalies and signals.{additional_sources}"
            f"{transcript_context}"
        )
    elif prompt_number == 4:
        return (
            f"Here are the raw signals from the Black Desk for {city}. "
            f"Run the full adversarial verification protocol on each signal.\n\n"
            f"--- BEGIN BLACK DESK OUTPUT ---\n{previous_output or '[No Black Desk output provided]'}\n--- END BLACK DESK OUTPUT ---"
        )
    elif prompt_number == 5:
        return (
            f"Run the five-part integrity audit on the following story package.\n\n"
            f"IMPORTANT: If the input below does not contain actual news stories (e.g., it "
            f"contains only questions, meta-commentary, or a refusal to write), state "
            f"'NO STORIES TO AUDIT — the previous stage did not produce stories' and stop. "
            f"Do not write an essay about the problem. Just report it in one sentence.\n\n"
            f"--- BEGIN STORY PACKAGE ---\n{previous_output or '[No story package provided]'}\n--- END STORY PACKAGE ---"
        )
    elif prompt_number == 6:
        if previous_output:
            return (
                f"Review the following civic stories for First Amendment and public records law issues.\n\n"
                f"--- BEGIN STORY ---\n{previous_output}\n--- END STORY ---"
            )
        if topic:
            return (
                f"Review the following topic for First Amendment and public records law issues "
                f"relevant to civic journalism:\n\nTOPIC: {topic}"
            )
        return "Provide a general overview of First Amendment protections for civic media operations."
    elif prompt_number == 7:
        return (
            f"Rewrite the following civic stories as community-ready news that residents can "
            f"paste into newsletters, email to council members, or share in neighborhood groups.\n\n"
            f"--- BEGIN STORY ---\n{previous_output or '[No story provided]'}\n--- END STORY ---"
        )
    elif prompt_number == 8:
        return (
            f"Run the Source Integrity Protocol check on the following story.\n\n"
            f"--- BEGIN STORY ---\n{previous_output or '[No story provided]'}\n--- END STORY ---"
        )
    elif prompt_number == 9:
        topic_text = topic or "general civic developments"
        # When chained from investigation pipeline (3→9), use topic leads as assignments
        investigation_context = ""
        investigation_mode = False
        if previous_output and not topic:
            investigation_mode = True
            investigation_context = (
                f"\n\nYou have {previous_output.count('TOPIC ')} story assignments below. "
                f"For each topic, use web_search to research it, then write a community "
                f"news story (400-800 words). Include source URLs for every claim.\n\n"
                f"After all stories, rank them by your confidence in your own sourcing:\n"
                f"- HIGH CONFIDENCE: Multiple independent sources confirmed the story\n"
                f"- MODERATE CONFIDENCE: Some corroboration but gaps remain\n"
                f"- LOW CONFIDENCE: Limited sourcing — worth reporting but thin\n\n"
                f"Put highest-confidence stories first.\n\n"
                f"--- STORY ASSIGNMENTS ---\n{previous_output}\n--- END ASSIGNMENTS ---\n\n"
            )
            topic_text = "the story assignments listed below"
        # In investigation mode, skip transcript injection — transcripts are raw
        # meeting footage that can cause the model to second-guess its assignments
        tx_ctx = "" if investigation_mode else transcript_context
        return (
            f"Research and write community-ready civic news stories for {city}.\n\n"
            f"TOPIC: {topic_text}\n\n"
            f"You MUST use web_search to find REAL, CURRENT information. Do NOT simulate or fabricate any facts. "
            f"Search for real sources, real quotes, and real data points.{additional_sources}"
            f"{investigation_context}"
            f"{tx_ctx}"
        )
    else:
        return previous_output or "Run this prompt."


class PipelineRunner:
    """Orchestrates running prompts through the Claude API."""

    # Prompts that benefit from YouTube meeting transcripts
    TRANSCRIPT_STAGES = {1, 2, 3, 9}

    def __init__(self, config):
        self.config = config
        self.outputs = {}  # {prompt_number: output_text}
        self._cancel = False
        self._client = None
        self._transcript_cache = None  # Fetched once per pipeline run

    def _get_client(self):
        """Lazily initialize the Anthropic client."""
        if self._client is None:
            try:
                from anthropic import Anthropic
            except ImportError:
                raise ImportError(
                    "The 'anthropic' package is required. Install it with:\n"
                    "  pip install anthropic"
                )
            api_key = self.config.get("api_key", "")
            if not api_key:
                raise ValueError("No API key configured. Please enter your Claude API key in Settings.")
            self._client = Anthropic(api_key=api_key)
        return self._client

    def cancel(self):
        """Signal the pipeline to stop after the current stage."""
        self._cancel = True

    def _get_transcript_context(self, on_progress=None):
        """
        Fetch YouTube meeting transcripts (cached per pipeline run).
        Checks ALL YouTube channels: the primary youtube_channel_url AND
        any YouTube URLs found in additional_sources.
        Returns formatted text to inject into user messages, or empty string.
        """
        if self._transcript_cache is not None:
            return self._transcript_cache

        try:
            from youtube_transcripts import (
                fetch_recent_meeting_transcripts,
                format_transcripts_for_prompt,
            )
        except ImportError:
            logger.warning(
                "youtube_transcripts module not available. "
                "Skipping transcript fetching."
            )
            self._transcript_cache = ""
            return ""

        # Collect ALL YouTube channel URLs from config
        yt_channels = []
        primary = self.config.get("sources", {}).get("youtube_channel_url", "")
        if primary:
            yt_channels.append(primary)
        for src in self.config.get("additional_sources", []):
            url = src.get("url", "")
            if "youtube.com" in url.lower():
                if url not in yt_channels:
                    yt_channels.append(url)

        if not yt_channels:
            self._transcript_cache = ""
            return ""

        if on_progress:
            on_progress(f"Fetching YouTube transcripts from {len(yt_channels)} channel(s)...")

        all_transcripts = []
        # Track which normalized titles we've collected WITH a transcript.
        # If the first copy has no transcript, keep looking on other channels.
        seen_with_transcript = set()
        seen_without_transcript = {}  # norm -> index in all_transcripts

        for channel_url in yt_channels:
            try:
                if on_progress:
                    on_progress(f"Checking YouTube channel: {channel_url}...")
                results = fetch_recent_meeting_transcripts(
                    channel_url, days=7, max_videos=3
                )
                for t in results:
                    norm = self._normalize_meeting_title(t.get("title", ""))
                    has_text = bool(t.get("transcript"))

                    # Already have this meeting WITH a transcript → skip
                    if norm in seen_with_transcript:
                        logger.info(
                            f"Skipping duplicate (already have transcript): "
                            f"{t['title']!r}"
                        )
                        if on_progress:
                            on_progress(
                                f"Skipped duplicate: {t['title'][:60]}"
                            )
                        continue

                    # We saw it before but WITHOUT a transcript
                    if norm in seen_without_transcript:
                        if has_text:
                            # This copy has one — replace the empty entry
                            idx = seen_without_transcript.pop(norm)
                            old_title = all_transcripts[idx]["title"]
                            logger.info(
                                f"Replacing transcript-less {old_title!r} "
                                f"with {t['title']!r} (has transcript)"
                            )
                            if on_progress:
                                on_progress(
                                    f"Found transcript on alt channel: "
                                    f"{t['title'][:60]}"
                                )
                            all_transcripts[idx] = t
                            seen_with_transcript.add(norm)
                        else:
                            # Neither copy has a transcript — skip
                            logger.info(
                                f"Skipping duplicate (neither has transcript): "
                                f"{t['title']!r}"
                            )
                            continue
                    else:
                        # First time seeing this meeting
                        idx = len(all_transcripts)
                        all_transcripts.append(t)
                        if has_text:
                            seen_with_transcript.add(norm)
                        else:
                            seen_without_transcript[norm] = idx
            except Exception as e:
                logger.error(f"Error fetching transcripts from {channel_url}: {e}")
                if on_progress:
                    on_progress(f"YouTube fetch failed for {channel_url}: {e}")

        self._transcript_cache = format_transcripts_for_prompt(all_transcripts)
        count = sum(1 for t in all_transcripts if t.get("transcript"))
        total = len(all_transcripts)
        if on_progress:
            if count:
                titles = [t["title"][:50] for t in all_transcripts if t.get("transcript")]
                on_progress(
                    f"Fetched {count} transcript(s) from {total} video(s): "
                    f"{', '.join(titles)}"
                )
            elif total > 0:
                on_progress(
                    f"Found {total} meeting video(s) but no transcripts available"
                )
            else:
                on_progress(f"No recent meeting videos found on {len(yt_channels)} YouTube channel(s)")

        return self._transcript_cache

    @staticmethod
    def _normalize_meeting_title(title: str) -> str:
        """
        Normalize a meeting video title for deduplication.

        Strips date format differences, extra whitespace, and punctuation so
        that 'City Council Regular Session - 03/10/2026' matches
        'City Council Regular Session - March 10, 2026'.
        """
        t = title.lower().strip()
        # Remove common date patterns:
        #   "March 10, 2026"  "03/10/2026"  "03-10-2026"  "2026-03-10"
        t = re.sub(
            r"(january|february|march|april|may|june|july|august|september|"
            r"october|november|december)\s+\d{1,2},?\s*\d{4}", "", t
        )
        # ISO format FIRST (2026-03-10) so MM/DD/YYYY doesn't partially eat it
        t = re.sub(r"\d{4}[/\-]\d{1,2}[/\-]\d{1,2}", "", t)
        t = re.sub(r"\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}", "", t)
        # Remove punctuation and extra whitespace
        t = re.sub(r"[^a-z0-9 ]", " ", t)
        t = re.sub(r"\s+", " ", t).strip()
        return t

    def run_single_stage(self, prompt_number, previous_output=None, topic=None,
                         on_progress=None, on_token=None):
        """
        Run a single pipeline stage.

        Args:
            prompt_number: Which prompt (1-9) to run
            previous_output: Output from the previous stage (if chaining)
            topic: Topic text (for prompts 6 and 9)
            on_progress: Callback(status_text) for progress updates
            on_token: Callback(token_text) for streaming tokens

        Returns:
            The full response text from Claude
        """
        self._cancel = False

        if on_progress:
            on_progress(f"Loading Prompt {prompt_number:02d}: {PROMPT_NAMES.get(prompt_number, '?')}...")

        # Load and prepare the prompt
        prompt_text = load_prompt_text(prompt_number)
        prompt_text = substitute_variables(prompt_text, self.config)

        # Fetch YouTube transcripts for stages that use them
        transcript_context = ""
        if prompt_number in self.TRANSCRIPT_STAGES:
            transcript_context = self._get_transcript_context(on_progress=on_progress)

        # Build the user message
        user_message = build_user_message(
            prompt_number,
            previous_output=previous_output,
            topic=topic,
            config=self.config,
            transcript_context=transcript_context,
        )

        client = self._get_client()

        # Use cheaper research model for web-search-heavy prompts (1, 3, 9)
        RESEARCH_STAGES = {1, 3, 9}
        research_model = self.config.get("research_model", "")
        if prompt_number in RESEARCH_STAGES and research_model:
            model = research_model
        else:
            model = self.config.get("model", "claude-haiku-4-5-20251001")

        if on_progress:
            on_progress(f"Sending to Claude ({model})...")

        # Only enable web search for stages that need live data.
        # Stages 2, 4, 5, 7, 8 process text from prior stages — no search needed.
        # Stage 6 only needs web search when NOT processing pasted content.
        WEB_SEARCH_STAGES = {1, 3, 9}
        if prompt_number == 6 and not previous_output:
            WEB_SEARCH_STAGES.add(6)
        if prompt_number in WEB_SEARCH_STAGES:
            # Stage 9 needs more searches when covering multiple topics
            # (investigation lane feeds it 5-8 topics to research independently)
            search_limit = 20 if prompt_number == 9 else 10
            tools = [
                {
                    "type": "web_search_20250305",
                    "name": "web_search",
                    "max_uses": search_limit,
                }
            ]
        else:
            tools = []

        # Use streaming for real-time output, with retry on rate limits
        MAX_RETRIES = 3
        RETRY_WAIT = [60, 90, 120]  # seconds — must exceed 1-min rate window
        full_response = []

        for attempt in range(1, MAX_RETRIES + 1):
            full_response = []
            try:
                with client.messages.stream(
                    model=model,
                    max_tokens=16384,
                    system=prompt_text,
                    messages=[{"role": "user", "content": user_message}],
                    tools=tools,
                ) as stream:
                    _cite_buf = ""  # Buffer for incomplete citation tags
                    for text in stream.text_stream:
                        if self._cancel:
                            break
                        full_response.append(text)
                        # Strip <cite> and </cite> tags from streaming display
                        if on_token:
                            _cite_buf += text
                            # If we might be mid-tag, hold the buffer
                            if "<" in _cite_buf and ">" not in _cite_buf:
                                continue  # Wait for closing >
                            # Strip complete citation tags and flush
                            clean = re.sub(
                                r'</?cite[^>]*>', '', _cite_buf
                            )
                            if clean:
                                on_token(clean)
                            _cite_buf = ""
                    # Flush any remaining buffer
                    if on_token and _cite_buf:
                        clean = re.sub(r'</?cite[^>]*>', '', _cite_buf)
                        if clean:
                            on_token(clean)
                break  # Success — exit retry loop

            except Exception as e:
                error_str = str(e).lower()
                is_rate_limit = (
                    "429" in str(e)
                    or "rate" in error_str
                    or "overloaded" in error_str
                    or "rate_limit" in error_str
                    or getattr(e, "status_code", None) == 429
                )
                logger.warning(
                    f"API error (attempt {attempt}/{MAX_RETRIES}): "
                    f"type={type(e).__name__}, rate_limit={is_rate_limit}, msg={e}"
                )
                if is_rate_limit and attempt < MAX_RETRIES:
                    wait_secs = RETRY_WAIT[attempt - 1]
                    if on_progress:
                        on_progress(
                            f"Rate limited — waiting {wait_secs}s before retry "
                            f"(attempt {attempt}/{MAX_RETRIES})..."
                        )
                    time.sleep(wait_secs)
                else:
                    raise  # Non-rate-limit error or final attempt

        output = "".join(full_response)
        # Strip API citation tags (e.g. <cite index="42-1,42-2">) from stored output
        output = re.sub(r'</?cite[^>]*>', '', output)
        self.outputs[prompt_number] = output

        if on_progress:
            if self._cancel:
                on_progress(f"Stage {prompt_number:02d} cancelled.")
            else:
                on_progress(f"Stage {prompt_number:02d} complete.")

        return output

    def run_lane(self, lane_key, topic=None, initial_input=None,
                 on_stage_start=None, on_progress=None, on_token=None,
                 on_stage_complete=None, on_lane_complete=None,
                 wait_for_approval=None):
        """
        Run a full pipeline lane.

        Args:
            lane_key: Key from LANES dict
            topic: Topic text (for ad-hoc lanes)
            initial_input: Pre-existing text to feed into the first stage
                           (e.g. story drafts for the story_polish lane)
            on_stage_start: Callback(stage_number, stage_name)
            on_progress: Callback(status_text)
            on_token: Callback(token_text) for streaming
            on_stage_complete: Callback(stage_number, output_text)
            on_lane_complete: Callback(all_outputs_dict)
            wait_for_approval: Callback(stage_number) -> bool (for step-by-step mode)
        """
        lane = LANES.get(lane_key)
        if not lane:
            raise ValueError(f"Unknown lane: {lane_key}")

        stages = lane["stages"]
        if not stages:
            raise ValueError("Single-prompt lane requires specifying a stage.")

        previous_output = initial_input  # None for most lanes; user text for story_polish
        INTER_STAGE_DELAY = 15  # seconds between stages to avoid rate limits

        # Some stages need output from a specific earlier stage, not just
        # the immediately preceding one.  e.g. Stage 7 (translator) needs
        # the stories from Stage 2, not the audit from Stage 5.
        INPUT_SOURCE = {
            7: 2,   # Community News Writer reads Story Expansion output
        }

        for i, stage_num in enumerate(stages):
            if self._cancel:
                break

            # Pause between stages so the rate-limit window resets
            if i > 0:
                if on_progress:
                    on_progress(f"Waiting {INTER_STAGE_DELAY}s between stages to avoid rate limits...")
                time.sleep(INTER_STAGE_DELAY)

            if on_stage_start:
                on_stage_start(stage_num, PROMPT_NAMES.get(stage_num, "Unknown"))

            # For ad-hoc story, first stage gets the topic
            stage_topic = topic if (stage_num == 9 or stage_num == 6) else None

            # Pick the right input: explicit source stage, or the previous stage
            if stage_num in INPUT_SOURCE and INPUT_SOURCE[stage_num] in self.outputs:
                stage_input = self.outputs[INPUT_SOURCE[stage_num]]
            else:
                stage_input = previous_output

            # Extract clean topic briefs when signal desk output feeds into Stage 9
            # Full analytical output from Stage 3 or 4 can cause Stage 9 to
            # refuse to write.  Pass only the topic names and source URLs.
            if stage_num == 9 and stage_input and i > 0:
                prev_stage = stages[i - 1] if i > 0 else None
                if prev_stage in (3, 4):
                    _before_len = len(stage_input)
                    stage_input = _extract_topic_briefs(
                        stage_input,
                        city_name=self.config.get("city_name", "")
                    )
                    _after_len = len(stage_input)
                    logger.info(f"Topic brief extraction: {_before_len} -> {_after_len} chars")
                    if on_progress:
                        on_progress(f"Extracted topic briefs: {_before_len} → {_after_len} chars")

            output = self.run_single_stage(
                prompt_number=stage_num,
                previous_output=stage_input,
                topic=stage_topic,
                on_progress=on_progress,
                on_token=on_token,
            )

            # Strip research preamble from Stage 9 output before it reaches Stage 5
            # Stage 9 narrates its search process before writing stories.
            # Stage 5 sees that preamble and thinks there are no stories.
            if stage_num == 9 and output:
                output = _strip_stage9_preamble(output)

            if on_stage_complete:
                on_stage_complete(stage_num, output)

            previous_output = output

            # Step-by-step: wait for user approval before next stage
            if wait_for_approval and i < len(stages) - 1:
                if not wait_for_approval(stage_num):
                    if on_progress:
                        on_progress("Pipeline paused by user.")
                    break

        if on_lane_complete and not self._cancel:
            on_lane_complete(self.outputs)

    def run_lane_async(self, lane_key, **kwargs):
        """Run a lane in a background thread."""
        thread = threading.Thread(
            target=self.run_lane,
            args=(lane_key,),
            kwargs=kwargs,
            daemon=True,
        )
        thread.start()
        return thread

    def save_outputs(self, directory=None):
        """Save all stage outputs as .md and .docx files."""
        if directory is None:
            directory = self.config.get("output_directory", "")
        if not directory:
            directory = "."

        output_dir = Path(directory)
        output_dir.mkdir(parents=True, exist_ok=True)

        today = datetime.now().strftime("%Y-%m-%d")
        saved_files = []

        for prompt_num, output in self.outputs.items():
            base = f"{today}_stage-{prompt_num:02d}_{PROMPT_FILES[prompt_num]}"

            # Save .md (raw text)
            md_path = output_dir / base
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(output)
            saved_files.append(str(md_path))

            # Save .docx (formatted Word document)
            docx_base = base.rsplit(".", 1)[0] + ".docx"
            docx_path = output_dir / docx_base
            try:
                _markdown_to_docx(output, docx_path)
                saved_files.append(str(docx_path))
            except ImportError:
                logger.warning(
                    "python-docx not installed — .docx export skipped. "
                    "Install with: pip install python-docx"
                )
            except Exception as e:
                logger.warning(f"Could not create .docx for stage {prompt_num}: {e}")

        return saved_files


def _markdown_to_docx(markdown_text, output_path):
    """Convert markdown-formatted text to a styled Word document.

    Handles: # H1, ## H2, ### H3, **bold** inline, --- horizontal rules.
    Requires python-docx.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Configure heading styles
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Segoe UI"
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0x67, 0x50, 0xA4)  # M3 primary purple

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped) or re.match(r"^─{3,}$", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCA, 0xC4, 0xD0)
            i += 1
            continue

        # Headings
        h_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2)
            # Strip bold markers from heading text
            heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # Empty line — skip (paragraph spacing handles gaps)
        if not stripped:
            i += 1
            continue

        # Normal paragraph — handle **bold** inline
        p = doc.add_paragraph()
        _add_runs_with_bold(p, stripped)
        i += 1

    doc.save(str(output_path))


def _add_runs_with_bold(paragraph, text):
    """Parse a line of text for **bold** markers and add styled runs."""
    from docx.shared import Pt

    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
